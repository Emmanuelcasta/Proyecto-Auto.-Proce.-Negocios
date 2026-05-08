"""
Generación de comprobante de nómina en formato DOCX.
"""

from io import BytesIO
from decimal import Decimal

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.models.nomina import Nomina


def _format_cop(valor: Decimal) -> str:
    n = f"{valor:,.2f}"
    # Formato es-CO manual: 1.234.567,89
    return "$ " + n.replace(",", "X").replace(".", ",").replace("X", ".")


def _add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _add_kv(document: Document, key: str, value: str) -> None:
    p = document.add_paragraph()
    k = p.add_run(f"{key}: ")
    k.bold = True
    p.add_run(value)


def _tabla_detalles(document: Document, titulo: str, detalles: list) -> None:
    if not detalles:
        return

    document.add_paragraph()
    t = document.add_paragraph()
    r = t.add_run(titulo)
    r.bold = True

    table = document.add_table(rows=1, cols=4)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Concepto"
    hdr[1].text = "Horas"
    hdr[2].text = "Porcentaje"
    hdr[3].text = "Valor"

    for d in detalles:
        row = table.add_row().cells
        row[0].text = d.concepto
        row[1].text = f"{d.horas:.2f}" if d.horas is not None else "-"
        row[2].text = f"{(d.porcentaje * Decimal('100')):.2f}%" if d.porcentaje is not None else "-"
        row[3].text = _format_cop(d.valor)


def generar_comprobante_docx(nomina: Nomina) -> bytes:
    """
    Genera un comprobante de nómina en DOCX y retorna bytes.
    """
    document = Document()

    _add_heading(document, "Comprobante de Nómina")
    _add_kv(document, "Empleado", nomina.empleado.nombre if nomina.empleado else f"ID {nomina.empleado_id}")
    _add_kv(document, "Periodo", f"{nomina.fecha_inicio} a {nomina.fecha_fin}")
    _add_kv(document, "Estado", nomina.estado)
    _add_kv(document, "Creado en", str(nomina.creado_en))
    if nomina.aprobado_en:
        _add_kv(document, "Aprobado en", str(nomina.aprobado_en))

    document.add_paragraph()
    _add_kv(document, "Días hábiles quincena", str(nomina.dias_habiles_quincena))
    _add_kv(document, "Días hábiles mes", str(nomina.dias_habiles_mes))
    _add_kv(document, "Umbral horas", f"{nomina.umbral_horas:.2f}")
    _add_kv(document, "Total horas trabajadas", f"{nomina.total_horas_trabajadas:.2f}")

    detalles = list(nomina.detalles or [])
    devengados = [d for d in detalles if d.categoria == "DEVENGADO"]
    deducciones = [d for d in detalles if d.categoria == "DEDUCCION"]
    empleador = [d for d in detalles if d.categoria == "EMPLEADOR"]

    _tabla_detalles(document, "Devengados", devengados)
    _tabla_detalles(document, "Deducciones", deducciones)
    _tabla_detalles(document, "Aportes Empleador (informativo)", empleador)

    document.add_paragraph()
    _add_kv(document, "Total devengado", _format_cop(nomina.total_devengado))
    _add_kv(document, "Total deducciones", _format_cop(nomina.total_deducciones))
    _add_kv(document, "Neto a pagar", _format_cop(nomina.neto_pagar))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
